"""Export combined clinical summaries to a Word document."""

from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.section import WD_ORIENT

from config.settings import settings


def _set_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.size = Pt(9)
    style.paragraph_format.space_after = Pt(2)
    style.paragraph_format.space_before = Pt(0)


def export_docx(
    summaries: list[tuple[str, str]],
    *,
    doctor_code: str = "",
    filename: str | None = None,
) -> Path:
    """Create a Word doc from a list of (chart_no, markdown) summaries."""
    doc = Document()

    # Landscape, narrow margins
    section = doc.sections[0]
    new_w, new_h = section.page_height, section.page_width
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = new_w
    section.page_height = new_h
    section.top_margin = Pt(36)
    section.bottom_margin = Pt(36)
    section.left_margin = Pt(36)
    section.right_margin = Pt(36)

    # Header
    header_para = section.header.paragraphs[0]
    run = header_para.add_run(
        f"日期: {datetime.now().strftime('%Y-%m-%d')}  醫師: {doctor_code}"
    )
    run.font.size = Pt(7)

    _set_style(doc)

    for i, (chart_no, md) in enumerate(summaries):
        if i > 0:
            doc.add_page_break()

        # Convert markdown to basic Word paragraphs
        for line in md.split("\n"):
            line = line.rstrip()
            if line.startswith("# "):
                p = doc.add_heading(line[2:], level=1)
                for run in p.runs:
                    run.font.size = Pt(12)
            elif line.startswith("## "):
                p = doc.add_heading(line[3:], level=2)
                for run in p.runs:
                    run.font.size = Pt(10)
            elif line.startswith("### "):
                p = doc.add_heading(line[4:], level=3)
                for run in p.runs:
                    run.font.size = Pt(9)
            elif line.startswith("- ") or line.startswith("* "):
                doc.add_paragraph(line[2:], style="List Bullet")
            elif line.strip():
                doc.add_paragraph(line)

    # Save
    out_dir = Path(settings.output_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    if filename is None:
        filename = f"{datetime.now().strftime('%Y%m%d')}_{doctor_code}_summaries.docx"
    path = out_dir / filename
    doc.save(str(path))
    return path
